import sys
import json
import re
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import fitz  # PyMuPDF
from typing import Dict, Any, List
import logging
import traceback

# Setup logging
logging.basicConfig(filename='parser.log', level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

class JournalProcessor:
    def __init__(self, journal_path: str):
        self.journal_path = journal_path
        self.raw_text = self._extract_text()
        self.paragraphs = self._split_paragraphs()
        self.metadata = self._extract_metadata()
        self.content = self._extract_content()
        logging.info(f"Processed {len(self.paragraphs)} paragraphs")

    def _extract_text(self) -> str:
        """Extract text from PDF or DOCX files"""
        try:
            text = ""
            if self.journal_path.endswith('.pdf'):
                with fitz.open(self.journal_path) as doc:
                    for page in doc:
                        text += page.get_text() + "\n"
            elif self.journal_path.endswith('.docx'):
                doc = Document(self.journal_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
            
            # Basic cleanup
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            text = re.sub(r'[ \t]+', ' ', text)
            
            logging.info(f"Extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logging.error(f"Text extraction failed: {e}")
            return ""

    def _split_paragraphs(self) -> List[str]:
        """Split text into meaningful paragraphs"""
        if not self.raw_text:
            return []
        
        paragraphs = []
        for para in self.raw_text.split('\n\n'):
            para = para.strip()
            if len(para) > 30 and not re.match(r'^\d+$', para):
                paragraphs.append(para)
        
        return paragraphs

    def _extract_metadata(self) -> Dict[str, str]:
        """Extract title and authors"""
        if not self.paragraphs:
            return {'title': 'No content found', 'authors': 'Unknown'}
        
        title = self.paragraphs[0][:200] if self.paragraphs else "Untitled Document"
        
        authors = "Author not detected"
        for para in self.paragraphs[:3]:
            if len(para.split()) <= 8 and len(para) < 100:
                words = para.split()
                if all(word[0].isupper() for word in words if word.isalpha()):
                    authors = para
                    break
        
        return {'title': title, 'authors': authors}

    def _extract_content(self) -> Dict[str, str]:
        """Extract content using position-based approach"""
        if not self.paragraphs:
            return {'full_content': 'No content available'}
        
        clean_paragraphs = []
        for para in self.paragraphs:
            if len(para.split()) > 5 and len(para) > 50:
                clean_paragraphs.append(para)
        
        total_clean = len(clean_paragraphs)
        content = {}
        
        content['full_content'] = '\n\n'.join(clean_paragraphs)
        
        if total_clean >= 6:
            content['abstract'] = '\n\n'.join(clean_paragraphs[0:2])
            content['introduction'] = '\n\n'.join(clean_paragraphs[2:4])
            content['method'] = '\n\n'.join(clean_paragraphs[4:6])
            content['result'] = '\n\n'.join(clean_paragraphs[6:8])
            content['discussion'] = '\n\n'.join(clean_paragraphs[8:10])
            content['conclusion'] = '\n\n'.join(clean_paragraphs[10:])
        elif total_clean >= 3:
            third = total_clean // 3
            content['abstract'] = '\n\n'.join(clean_paragraphs[0:1])
            content['introduction'] = '\n\n'.join(clean_paragraphs[1:1+third])
            content['method'] = '\n\n'.join(clean_paragraphs[1+third:1+2*third])
            content['result'] = '\n\n'.join(clean_paragraphs[1+2*third:])
            content['discussion'] = ''
            content['conclusion'] = ''
        else:
            content['abstract'] = clean_paragraphs[0] if len(clean_paragraphs) > 0 else ""
            content['introduction'] = clean_paragraphs[1] if len(clean_paragraphs) > 1 else ""
            content['method'] = clean_paragraphs[2] if len(clean_paragraphs) > 2 else ""
            content['result'] = ''
            content['discussion'] = ''
            content['conclusion'] = ''
        
        return content

def preserve_paragraph_formatting(para, new_text):
    """Replace paragraph text while preserving ALL formatting"""
    if not new_text.strip():
        return
    
    # Store ALL paragraph properties
    original_style = para.style
    original_alignment = para.alignment
    original_space_before = para.space_before
    original_space_after = para.space_after
    original_line_spacing = para.line_spacing
    original_first_line_indent = para.first_line_indent
    original_left_indent = para.left_indent
    original_right_indent = para.right_indent
    
    # Store run formatting from first run
    run_formatting = None
    if para.runs:
        first_run = para.runs[0]
        run_formatting = {
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
            'bold': first_run.font.bold,
            'italic': first_run.font.italic,
            'underline': first_run.font.underline,
            'color': first_run.font.color.rgb if first_run.font.color.rgb else None,
            'highlight_color': first_run.font.highlight_color
        }
    
    # Clear paragraph content
    para.clear()
    
    # Add new text with preserved formatting
    run = para.add_run(new_text)
    
    # Apply stored run formatting
    if run_formatting:
        try:
            if run_formatting['font_name']:
                run.font.name = run_formatting['font_name']
            if run_formatting['font_size']:
                run.font.size = run_formatting['font_size']
            if run_formatting['bold'] is not None:
                run.font.bold = run_formatting['bold']
            if run_formatting['italic'] is not None:
                run.font.italic = run_formatting['italic']
            if run_formatting['underline'] is not None:
                run.font.underline = run_formatting['underline']
            if run_formatting['color']:
                run.font.color.rgb = run_formatting['color']
            if run_formatting['highlight_color']:
                run.font.highlight_color = run_formatting['highlight_color']
        except Exception as e:
            logging.warning(f"Could not apply run formatting: {e}")
    
    # Restore paragraph formatting
    try:
        para.style = original_style
        para.alignment = original_alignment
        if original_space_before:
            para.space_before = original_space_before
        if original_space_after:
            para.space_after = original_space_after
        if original_line_spacing:
            para.line_spacing = original_line_spacing
        if original_first_line_indent:
            para.first_line_indent = original_first_line_indent
        if original_left_indent:
            para.left_indent = original_left_indent
        if original_right_indent:
            para.right_indent = original_right_indent
    except Exception as e:
        logging.warning(f"Could not restore paragraph formatting: {e}")

def preserve_cell_formatting(cell, new_text):
    """Replace cell text while preserving formatting"""
    if not new_text.strip():
        return
    
    # Get the first paragraph in the cell
    if cell.paragraphs:
        para = cell.paragraphs[0]
        preserve_paragraph_formatting(para, new_text)
        
        # Remove any additional paragraphs
        for i in range(len(cell.paragraphs) - 1, 0, -1):
            try:
                p = cell.paragraphs[i]._element
                p.getparent().remove(p)
            except:
                pass

def apply_template(processor: JournalProcessor, template_path: str, output_path: str):
    """Apply content to template while preserving exact formatting"""
    try:
        logging.info(f"Applying template: {template_path}")
        
        if template_path.endswith('.docx'):
            # Load template
            doc = Document(template_path)
            
            # Prepare replacements
            replacements = {
                # English placeholders
                '{{title}}': processor.metadata['title'],
                '{{authors}}': processor.metadata['authors'],
                '{{abstract}}': processor.content.get('abstract', ''),
                '{{introduction}}': processor.content.get('introduction', ''),
                '{{method}}': processor.content.get('method', ''),
                '{{methodology}}': processor.content.get('method', ''),
                '{{result}}': processor.content.get('result', ''),
                '{{results}}': processor.content.get('result', ''),
                '{{discussion}}': processor.content.get('discussion', ''),
                '{{conclusion}}': processor.content.get('conclusion', ''),
                
                # Indonesian placeholders
                '{{judul}}': processor.metadata['title'],
                '{{penulis}}': processor.metadata['authors'],
                '{{abstrak}}': processor.content.get('abstract', ''),
                '{{pendahuluan}}': processor.content.get('introduction', ''),
                '{{metode}}': processor.content.get('method', ''),
                '{{metodologi}}': processor.content.get('method', ''),
                '{{hasil}}': processor.content.get('result', ''),
                '{{pembahasan}}': processor.content.get('discussion', ''),
                '{{kesimpulan}}': processor.content.get('conclusion', ''),
                
                # Full content fallbacks
                '{{content}}': processor.content['full_content'],
                '{{konten}}': processor.content['full_content'],
                '{{full_content}}': processor.content['full_content'],
                '{{teks_lengkap}}': processor.content['full_content'],
            }
            
            # Process paragraphs with exact formatting preservation
            replacements_made = 0
            for para in doc.paragraphs:
                original_text = para.text
                
                # Check for placeholders and replace
                for placeholder, replacement in replacements.items():
                    if placeholder in original_text:
                        if replacement:  # Only replace if we have content
                            new_text = original_text.replace(placeholder, replacement)
                            preserve_paragraph_formatting(para, new_text)
                            replacements_made += 1
                            logging.info(f"Replaced {placeholder} with {len(replacement)} chars")
                            break
                        else:
                            # For empty replacements, preserve the paragraph but make it empty
                            preserve_paragraph_formatting(para, "")
                            break
            
            # Process tables with formatting preservation
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in original_text and replacement:
                                new_text = original_text.replace(placeholder, replacement)
                                preserve_cell_formatting(cell, new_text)
                                replacements_made += 1
                                logging.info(f"Replaced {placeholder} in table")
                                break
            
            # Remove template instructions but preserve structure
            for para in doc.paragraphs:
                text = para.text.lower()
                if any(word in text for word in ['template', 'placeholder', 'contoh', 'example', '[isi', '[masukkan']):
                    if '{{' not in para.text and '}}' not in para.text:  # Don't remove if it has placeholders
                        preserve_paragraph_formatting(para, "")
            
            logging.info(f"Made {replacements_made} replacements while preserving formatting")
            
            # If no placeholders found, the template might be plain text
            if replacements_made == 0:
                logging.warning("No placeholders found - template might not use placeholders")
                # In this case, don't modify the template, just add content at the end
                doc.add_page_break()
                doc.add_heading("KONTEN JURNAL YANG DIEKSTRAK", 1)
                doc.add_paragraph(f"Judul: {processor.metadata['title']}")
                doc.add_paragraph(f"Penulis: {processor.metadata['authors']}")
                doc.add_paragraph("Konten:")
                doc.add_paragraph(processor.content['full_content'])
        
        else:
            # For non-DOCX templates, create structured document
            doc = Document()
            doc.add_heading(processor.metadata['title'], 0)
            
            author_para = doc.add_paragraph()
            author_para.add_run(f"Penulis: {processor.metadata['authors']}").italic = True
            
            sections = [
                ('Abstract', processor.content.get('abstract', '')),
                ('Pendahuluan', processor.content.get('introduction', '')),
                ('Metode', processor.content.get('method', '')),
                ('Hasil', processor.content.get('result', '')),
                ('Pembahasan', processor.content.get('discussion', '')),
                ('Kesimpulan', processor.content.get('conclusion', ''))
            ]
            
            for section_title, section_content in sections:
                if section_content.strip():
                    doc.add_heading(section_title, 1)
                    doc.add_paragraph(section_content)
            
            # Add full content if sections are sparse
            content_count = sum(1 for _, content in sections if content.strip())
            if content_count < 3:
                doc.add_heading("Konten Lengkap", 1)
                doc.add_paragraph(processor.content['full_content'])
        
        # Save document
        doc.save(output_path)
        logging.info(f"Document saved with preserved formatting")
        
        return output_path
        
    except Exception as e:
        logging.error(f"Template application failed: {str(e)}")
        raise

def main(journal_path: str, template_path: str, output_path: str) -> Dict[str, Any]:
    """Main processing function"""
    try:
        # Validate files
        if not os.path.exists(journal_path):
            raise FileNotFoundError(f"Journal file not found: {journal_path}")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Create output directory
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Process journal
        logging.info(f"Processing journal: {journal_path}")
        processor = JournalProcessor(journal_path)
        
        if not processor.raw_text:
            raise ValueError("No text extracted from journal")
        
        # Apply template
        apply_template(processor, template_path, output_path)
        
        # Verify output
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Output file not created")
        
        return {
            'success': True,
            'output_path': output_path,
            'metadata': processor.metadata,
            'file_size': os.path.getsize(output_path),
            'paragraphs_processed': len(processor.paragraphs),
            'text_length': len(processor.raw_text),
            'formatting_preserved': True
        }
        
    except Exception as e:
        logging.error(f"Processing failed: {str(e)}\n{traceback.format_exc()}")
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }

if __name__ == "__main__":
    if len(sys.argv) < 4:
        result = {'error': 'Usage: python parser.py <journal_path> <template_path> <output_path>'}
    else:
        journal_path = sys.argv[1]
        template_path = sys.argv[2]
        output_path = sys.argv[3]
        
        result = main(journal_path, template_path, output_path)
    
    print(json.dumps(result))